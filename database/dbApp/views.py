import json
from django.http import HttpResponse
from django.shortcuts import render
from dbApp.models import Categories
import pyodbc
import csv
import openpyxl
from io import BytesIO
from docx import Document
from weasyprint import HTML
from django.template.loader import render_to_string

# Create your views here.

def Showcategories(request):
    categories = Categories.objects.all()
    dict = {'categories': categories}
    
    return render(request, 'dbApp/category.html', dict)

def rawSql(request):
    query = '''
        SELECT a.OrderID, a.OrderDate, b.CompanyName,
        c.ProductName, d.UnitPrice, d.Quantity,
        d.unitprice * d.Quantity as 'BillAmount'
        FROM orders 
        a inner join [order details] d on a.orderid = d.orderid 
        inner join
        customers b on a.customerid = b.customerid 
        inner join 
        products c on d.productid = c.productid
        WHERE a.orderid BETWEEN 10248 AND 10255
    '''
    conn = getConnection()  #connect to the database
    cursor = conn.cursor()  #Create a new cursor
    cursor.execute(query)   
    orders = cursor.fetchall()
    dict = {'orders': orders}
    return render(request, 'dbApp/showorder.html', dict)

def storedProcedureSQl(request):

    # GrandTotal = 0
    # runningTotal = 0
    # runningOrderTotal = 0

    # newOrders = []

    # for order in orders:
    #     runningTotal += order.BillAmount
    #     runningOrderTotal += order.BillAmount
    #     GrandTotal += order.BillAmount
    #     newOrders.append(pushData(order, runningTotal, runningOrderTotal))
   # dict = {'oder': order}
    #return render(request, 'dbApp/showorder.html', dict)

    procedureName = 'USP_GetAllOrder'
    conn = getConnection()  #connect to the database
    cursor = conn.cursor()  #Create a new cursor
    cursor.execute('{call USP_GetAllOrder}')   
    orders = cursor.fetchall()
    dict = {'orders': orders}
    #dict = {'orders': newOrders, 'GrandTotal': GrandTotal}
    return render(request, 'dbApp/showorder.html', dict)

# def pushData(order, runningTotal, runningOderTotal):

#     dataToPush = {
#         'OrderID': order.OrderID,
#         'OrderDate': order.OrderDate,
#         'CompanyName': order.CompanyName,
#         'ProductName': order.ProductName,
#         'UnitPrice': order.UniPrice,
#         'Quantity': order.Quantity,
#         'BillAmount': order.BillAmount,
#         'RunningTotal': runningTotal,
#         'RunningOrderTotal': runningOderTotal
#     }
#     return(dataToPush)

def storedProcedureWithoutOutputParameters(request):
    procedureName = 'USP_GetAllOrder'
    conn = getConnection()  #connect to the database
    cursor = conn.cursor()  #Create a new cursor
    count = 0
    #Procedure with paramaters
    cursor.execute('{call USP_GetOrdersCount(?)}', count)   
    orders = cursor.fetchval()

    #Procedure without paramaters
    cursor.execute('{call USP_GetAllOrder}')   
    orders = cursor.fetchall()
    dict = {'orders': orders, 'count': count}
    return render(request, 'dbApp/showorder.html', dict)


def getConnection():
    conn = pyodbc.connect('DRIVER=ODBC Driver 17 for SQL Server; Server=.;Database=northwind; Trusted_Connection=Yes;')
    #conn = pyodbc.connect('DRIVER= SQL Server Native Client 11.0; Server=.;Database=Northwind; Trusted_Connection=Yes;')
    return(conn)


#Exporting Data To CSV
def exporttoCSV(request):
    categories = Categories.objects.all()
    file_name = f'category_data.csv'
    response = HttpResponse(content_type= 'text/csv')
    response['Content-Disposition'] = f'attachment; filename= "{file_name}"'
    writer = csv.writer(response)

    writer.writerow(['Category ID', 'Category Name', 'Decription'])
    for category in categories:
        writer.writerow([category.categoryid, category.categoryname, category.description])

    return response

#Exporting Data To Json
def exporttoJSON(request):
    categories = Categories.objects.all()
    file_name = f'category_data.json'
    response = HttpResponse(content_type= 'text/json')
    response['Content-Disposition'] = f'attachment; filename= "{file_name}"'

    data = [
        {
            'categoryid': category.categoryid,
            'categoryname': category.categoryname,
            'description': category.description
        }
        for category in categories]
    json.dump(data, response)

    return response

#Exporting Data To Excel
def exportToXLS(request):
    categories = Categories.objects.all()
    file_name = f'category_data.xlsx'

    #Generate an Excel Workbook
    workbook = openpyxl.Workbook()
    worksheet = workbook.active

    #Add headers
    headers = ['Category ID', 'Category Name', 'Decription']
    worksheet.append(headers)

    #Add data
    for category in categories:
        worksheet.append([
            category.categoryid, category.categoryname, category.description
        ])

    #Save the workbook to a BytesID buffer
    buffer = BytesIO()
    workbook.save(buffer)
    buffer.seek(0)

    #Return the Excel file as the HTTP response
    response = HttpResponse(buffer.read(),
                            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'
    return response     

#Exporting Data To Word
def exportToWord(request):
    categories = Categories.objects.all()
    file_name = f'category_data.docx'

    #Generate a word document
    document = Document()

    #Add a table with headers
    table = document.add_table(rows=1, cols=3)
    table.style= 'TableGrid'
    hearder_row= table.rows[0].cells
    hearder_row[0].text= 'Category ID'
    hearder_row[1].text= 'Category Name'
    hearder_row[2].text= 'Decription' 
  
    #Bold the header
    for cell in hearder_row:
        cell.paragraphs[0].runs[0].font.bold = True

    #Add data to the table
    for category in categories:
        row = table.add_row().cells #create blank row
        row[0].text = str(category.categoryid)
        row[1].text = category.categoryname
        row[2].text = category.description

     #Save the word document to a BytesID buffer
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    #Return the Word file as the HTTP response
    response = HttpResponse(buffer.read(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{file_name}"'
    return response    


def exportPdf(request):
    categories = Categories.objects.all()
    file_name = f'category_data.pdf'

    response = HttpResponse(content_type= 'application/pdf')
    response['Content-Disposition']= f'attachment; filename="{file_name}"'

    dict = {'categories': categories}
    html_string = render_to_string('dbApp/showPdf.html', dict)
    html = HTML(string=html_string, base_url='')
    html.write_pdf(response)

    return response

#'DRIVER': 'SQL Server Native Client 11.0'
#DRIVER={ODBC Driver 17 for SQL Server}